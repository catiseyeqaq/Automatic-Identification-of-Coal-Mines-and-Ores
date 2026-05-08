from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BASE = Path(r"D:\ultralytics-yolov8")
MODEL_WEIGHT = r"D:\ultralytics-yolov8\runs\train\YOLO-GDL\weights\best.pt"


SECTIONS = [
    (
        "文档定位",
        [
            "本文件为煤矿履带煤块与矸石智能识别系统的当前技术路线说明，已根据项目目录、训练结果、模型配置、上位机脚本、Qoder 代码知识库以及论文实验章节重新同步。",
            "旧版技术手册主要描述“YOLOv8 + Gradio + 串口控制”的总体方案；当前项目已经形成以 YOLO-GDL 为最终算法模型的实现路线，本文档据此更新模型命名、权重路径、实验指标、软件流程和接口边界。",
        ],
    ),
    (
        "一、当前项目组成",
        [
            r"核心代码仓库位于 D:\ultralytics-yolov8，主体仍基于 Ultralytics YOLO 工程。项目新增或重点使用的文件包括：coal_gangue_app.py、trainv8.py、yolov8n.yaml、煤块和石块\data.yaml、ultralytics\cfg\models\v8\yolov8n-ghostneck-dwdown.yaml，以及 runs\train 下的训练结果。",
            r"Qoder 资料位于 .qoder\repowiki\zh，主要提供 Ultralytics 训练流程、目标检测任务、预测器、实时推理、Gradio 队列等代码级说明，可作为理解框架调用链的辅助资料；本系统的业务模型、权重路径和实验指标以当前工程文件和 runs 目录为准。",
        ],
    ),
    (
        "二、系统总体架构",
        [
            "系统采用“上位机 AI 视觉与 Web 控制台 + 下位机物理执行机构”的主从解耦架构。上位机负责图像采集、YOLO-GDL 推理、检测结果统计、语音播报、日志记录和 Web 交互；下位机负责履带启停、照明开关以及后续可扩展的剔除执行。",
            "上位机软件由 Python、Ultralytics YOLO、OpenCV、Gradio、pyttsx3 和 pyserial 组成。检测引擎封装在 CoalGangueDetector 中，支持图片、视频和摄像头实时流三类输入，并使用线程锁避免多线程同时占用 GPU 推理造成阻塞。",
            "下位机建议采用 STM32 或 ESP32。当前代码中已经实现 UART 串口 JSON 指令发送，并在未安装 pyserial 或硬件未连接时自动退化为日志模拟模式，便于本科设计演示和分阶段联调。",
        ],
    ),
    (
        "三、YOLO-GDL 算法技术路线",
        [
            "当前最终模型名称为 YOLO-GDL，来源于 GhostNeck、DWDown 和 Light-aware Augmentation 三个核心改进。G 表示使用 GhostNeck 轻量化 Neck 特征融合；D 表示使用 DWConv 完成 Neck 下采样；L 表示面向煤块和矸石颜色接近、光照变化明显的场景引入光照增强。",
            "YOLO-GDL 没有全量替换 Backbone，而是保留 YOLOv8n 的主干特征提取结构，将轻量化重点放在 Neck 区域。这样可以保留煤块与矸石细粒度纹理表达能力，同时减少特征融合和下采样阶段的冗余计算。",
            r"最终模型配置文件为 ultralytics\cfg\models\v8\yolov8n-ghostneck-dwdown.yaml。该配置将 Head/Neck 中的 C2f 替换为 C3Ghost，并将两处下采样 Conv 替换为 DWConv，检测头仍输出 P3、P4、P5 三尺度结果，类别数 nc=2。",
        ],
    ),
    (
        "四、数据集与训练配置",
        [
            r"数据配置文件为 煤块和石块\data.yaml，任务类别为 coal 与 gangue 两类。训练集路径为 train/images，验证集路径为 valid/images。",
            r"当前训练实验统一使用 1280 输入尺寸、batch size 8、epochs 100、patience 25、optimizer auto、GPU device 0 和 AMP 自动混合精度。消融实验保存在 runs\train 下，包含 YOLOv8n 基准、光照增强、PConv/C2f-Faster、GhostNeck、GhostDown、EMA、DWDown 等多组对比。",
            r"需要注意：根目录 trainv8.py 是早期 YOLOv8n 从头训练脚本，仍保留为基准训练入口；YOLO-GDL 的最终实验结果来自 runs\train\YOLO-GDL，正式说明和交付应优先引用 YOLO-GDL 权重。",
        ],
    ),
    (
        "五、实验结果与最终权重",
        [
            r"根据当前 runs\train 中 results.csv 的最佳指标，最终模型为 YOLO-GDL，对应权重路径为 D:\ultralytics-yolov8\runs\train\YOLO-GDL\weights\best.pt。",
            "YOLO-GDL 在验证集上的关键指标为 Precision=0.91394，Recall=0.92207，mAP50=0.97231，mAP50-95=0.84803。该结果高于当前目录中的 YOLO26n baseline，也优于普通 YOLOv8n 基准的综合轻量化表现。",
            "技术结论：YOLO-GDL 在保持检测精度提升的同时降低模型规模，适合作为本系统当前最终采用的煤矸石目标检测模型。后续论文、演示系统和部署说明均应统一使用 YOLO-GDL 命名与最终权重路径。",
        ],
    ),
    (
        "六、上位机软件实现",
        [
            "上位机入口脚本为 coal_gangue_app.py，功能包括图片检测、视频检测、摄像头实时检测、检测统计、语音播报、日志记录、串口连接管理、履带控制、照明控制和应急停止。",
            "Gradio Blocks 负责构建 Web 控制台，OpenCV 负责图像与视频帧处理，Ultralytics YOLO 负责推理，pyttsx3 负责本地语音播报，pyserial 负责 UART 串口通信。实时摄像头模式采用后台线程持续采集与推理，并通过轮询将最新帧、统计结果和截图画廊刷新到前端。",
            "当前脚本支持命令行参数或 COAL_MODEL_PATH 环境变量覆盖模型路径。若不指定路径，应将默认模型同步为 YOLO-GDL 的 best.pt，以避免误加载早期 YOLOv8n 基准权重。",
        ],
    ),
    (
        "七、通信协议与硬件边界",
        [
            "上位机与下位机之间采用 USB 转 TTL 串口通信，推荐波特率 115200bps，数据格式为 UTF-8 JSON 字符串，每条指令以换行结束。",
            '{"cmd":"belt","action":"start"}、{"cmd":"belt","action":"stop"}、{"cmd":"light","action":"on"}、{"cmd":"light","action":"off"} 是当前已经实现的基础控制指令。应急停止逻辑会同时发送履带停止和照明关闭指令。',
            "杂质剔除机构当前在旧版技术路线中作为硬件扩展目标描述，当前上位机代码已具备检测告警和控制接口基础，但自动剔除触发协议仍需硬件端确认具体动作、延时补偿和安全互锁后再固化。",
        ],
    ),
    (
        "八、与旧版 Word 技术手册的主要差异",
        [
            "旧版文档：核心算法写为“基于 YOLOv8 架构的目标检测模型”。当前实现：最终模型已明确为 YOLO-GDL，并有对应模型配置、消融实验和最终权重。",
            "旧版文档：侧重总体“感知 -> 决策 -> 执行”流程。当前实现：除了总体流程，还需要写明 GhostNeck、DWDown、Light-aware Augmentation、两类数据集、训练超参和验证指标。",
            "旧版文档：硬件执行包含剔除模块目标。当前实现：代码已实现履带、照明、语音、日志和 Web 控制台，剔除机构仍属于待与下位机联调确认的扩展执行项。",
            r"旧版文档：没有指定最终权重。当前实现：最终权重为 runs\train\YOLO-GDL\weights\best.pt，论文实验章节也已采用 YOLO-GDL 命名。",
        ],
    ),
    (
        "九、当前待同步事项",
        [
            r"建议将 coal_gangue_app.py 的默认模型路径从早期 coal-gangue-yolov8n 从头训练权重更新为 runs\train\YOLO-GDL\weights\best.pt，或在运行时始终通过命令行参数/COAL_MODEL_PATH 指定 YOLO-GDL 权重。",
            "建议后续新增 README 或运行说明，明确训练入口、推理入口、最终模型路径和串口硬件连接方式，避免论文模型、演示系统模型和实验目录不一致。",
            "若进入硬件联调阶段，需要补充自动剔除触发策略，包括检测到 gangue 后的延时、履带速度标定、执行机构动作时序、误触发保护和急停优先级。",
        ],
    ),
]


TABLES = {
    "一、当前项目组成": (
        "关键文件与作用",
        ["文件/目录", "当前作用"],
        [
            [
                r"ultralytics\cfg\models\v8\yolov8n-ghostneck-dwdown.yaml",
                "YOLO-GDL 最终模型结构配置：C3Ghost + DWConv Neck 下采样",
            ],
            [r"runs\train\YOLO-GDL\weights\best.pt", "当前最终改进模型权重"],
            [r"煤块和石块\data.yaml", "coal/gangue 两分类数据集配置"],
            ["coal_gangue_app.py", "Gradio 上位机应用，支持图片、视频、摄像头、语音、日志和串口控制"],
            ["trainv8.py", "早期 YOLOv8n 基准训练脚本，保留作对照"],
            [r".qoder\repowiki\zh", "Ultralytics、训练、推理、Gradio 等框架知识库资料"],
        ],
    ),
    "五、实验结果与最终权重": (
        "最终模型指标",
        ["模型", "Precision", "Recall", "mAP50", "mAP50-95", "权重路径"],
        [["YOLO-GDL", "0.91394", "0.92207", "0.97231", "0.84803", MODEL_WEIGHT]],
    ),
    "九、当前待同步事项": (
        "当前运行建议",
        ["场景", "建议"],
        [
            ["论文与答辩", "统一称为 YOLO-GDL，引用最终权重和消融实验结果"],
            ["演示系统", "运行 coal_gangue_app.py 时显式传入 YOLO-GDL best.pt 或设置 COAL_MODEL_PATH"],
            ["后续训练", "以 yolov8n-ghostneck-dwdown.yaml 为最终模型配置继续微调"],
            ["硬件联调", "先使用串口模拟模式验证 JSON 指令，再接入 STM32/ESP32"],
        ],
    ),
}


def font_run(run, size: float = 10.5, bold: bool = False, color: tuple[int, int, int] | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    font_run(r, 9, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    font_run(r, 10.5)


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]]) -> None:
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
        shade(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def locate_target() -> Path:
    candidates = [
        p
        for p in BASE.glob("*.docx")
        if not p.name.startswith("YOLO-GDL") and "_backup_" not in p.name and "backup" not in p.name.lower()
    ]
    if not candidates:
        raise FileNotFoundError("No technical manual docx found")
    return candidates[0]


def main() -> None:
    target = locate_target()
    backup = BASE / f"tech_manual_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    backup.write_bytes(target.read_bytes())

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    for style_name, size in [("Normal", 10.5), ("Title", 18), ("Heading 1", 14), ("Heading 2", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_run(p.add_run("煤矿履带煤块与矸石智能识别系统"), 18, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_run(p.add_run("当前技术路线与 YOLO-GDL 模型更新说明书"), 13, False, (80, 80, 80))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_run(
        p.add_run(
            "更新依据：项目文件遍历、runs 训练结果、YOLO-GDL 论文实验章节、Qoder 框架文档 | 更新日期：2026-05-05"
        ),
        9,
        False,
        (100, 100, 100),
    )

    for title, paragraphs in SECTIONS:
        doc.add_heading(title, level=1)
        for paragraph in paragraphs:
            add_para(doc, paragraph)
        if title in TABLES:
            add_table(doc, *TABLES[title])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font_run(p.add_run(r"最终模型：YOLO-GDL | 权重：runs\train\YOLO-GDL\weights\best.pt"), 9, False, (90, 90, 90))

    doc.save(target)
    print(f"UPDATED={target}")
    print(f"BACKUP={backup}")


if __name__ == "__main__":
    main()
