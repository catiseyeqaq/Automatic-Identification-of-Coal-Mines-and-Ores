from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
PAPER = ROOT / "YOLO-GDL_煤矸石识别论文实验章节_.docx"
BACKUP = ROOT / "YOLO-GDL_煤矸石识别论文实验章节_备份_加入训练代码前.docx"


TRAIN_CODE = """from ultralytics import YOLO


def main():
    train_args = {
        "data": "煤块和石块/data.yaml",
        "epochs": 100,
        "imgsz": 1280,
        "batch": 8,
        "workers": 4,
        "device": 0,
        "amp": True,
        "patience": 25,
        "project": "runs/train",
        "name": "YOLO-GDL",
        "exist_ok": True,
        "pretrained": True,
        "optimizer": "auto",
        "seed": 0,
        "deterministic": True,
        "lr0": 0.01,
        "lrf": 0.01,
        "momentum": 0.937,
        "weight_decay": 0.0005,
        "warmup_epochs": 3.0,
        "close_mosaic": 10,
        "hsv_h": 0.01,
        "hsv_s": 0.35,
        "hsv_v": 0.65,
        "translate": 0.1,
        "scale": 0.5,
        "fliplr": 0.5,
        "mosaic": 1.0,
        "plots": True,
        "save": True,
    }

    model = YOLO("ultralytics/cfg/models/v8/yolov8n-ghostneck-dwdown.yaml")
    model.train(**train_args)


if __name__ == "__main__":
    main()"""


PARAM_ROWS = [
    (
        "模型配置",
        "ultralytics/cfg/models/v8/yolov8n-ghostneck-dwdown.yaml",
        "YOLO-GDL 最终结构，Neck 采用 C3Ghost，Neck 下采样采用 DWConv",
    ),
    ("数据集配置", "煤块和石块/data.yaml", "训练集 train/images，验证集 valid/images，类别为 coal 和 gangue"),
    ("训练轮次 epochs", "100", "保证小样本煤矸石数据充分收敛"),
    ("输入尺寸 imgsz", "1280", "保留履带图像中煤块、矸石目标的细节纹理"),
    ("批大小 batch", "8", "兼顾显存占用与梯度稳定性"),
    ("训练设备 device", "0", "使用第 0 块 GPU 训练"),
    ("数据加载 workers", "4", "Windows 环境下的数据加载进程数"),
    ("混合精度 amp", "True", "启用自动混合精度以提升训练效率"),
    ("早停 patience", "25", "验证指标连续 25 轮无提升时触发早停"),
    ("优化器 optimizer", "auto", "使用 Ultralytics 自动选择优化器"),
    ("初始学习率 lr0", "0.01", "训练初始学习率"),
    ("最终学习率系数 lrf", "0.01", "训练结束时学习率为 lr0 的 0.01 倍"),
    ("动量 momentum", "0.937", "SGD/优化器动量参数"),
    ("权重衰减 weight_decay", "0.0005", "抑制过拟合"),
    ("预热轮次 warmup_epochs", "3.0", "训练前期逐步升高学习率"),
    ("关闭 Mosaic 轮次 close_mosaic", "10", "最后 10 轮关闭 Mosaic，提升真实分布适应性"),
    ("HSV 增强", "hsv_h=0.01, hsv_s=0.35, hsv_v=0.65", "模拟现场光照、曝光和颜色波动"),
    ("几何增强", "translate=0.1, scale=0.5, fliplr=0.5", "提升目标位置、尺度和方向变化下的鲁棒性"),
    ("Mosaic 增强", "mosaic=1.0", "提高复杂场景和目标组合的泛化能力"),
    ("随机种子", "seed=0, deterministic=True", "增强实验复现性"),
    ("结果保存", "runs/train/YOLO-GDL", "保存权重、训练曲线、混淆矩阵和验证结果"),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def move_after(anchor, elements) -> None:
    current = anchor._p
    for element in elements:
        current.addnext(element)
        current = element


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(PAPER, BACKUP)

    doc = Document(PAPER)
    if any("训练代码与参数设置" in p.text for p in doc.paragraphs):
        print("training section already exists")
        return

    anchor = next(p for p in doc.paragraphs if "为验证各改进模块的有效性" in p.text)

    elements = []

    h = doc.add_paragraph(style="Heading 2")
    h.add_run("4.1 训练代码与参数设置")
    elements.append(h._p)

    p = doc.add_paragraph()
    p.add_run(
        "为保证实验结果可复现，本文在训练阶段固定模型配置文件、数据集划分、输入分辨率、训练轮次和随机种子。"
        "最终模型 YOLO-GDL 以 YOLOv8n 为基础，在 Neck 部分引入 GhostNeck 轻量化特征融合结构，"
        "并将 Neck 下采样卷积替换为 DWDown 深度可分离下采样模块；训练时结合 Light-aware 光照增强策略。"
    )
    elements.append(p._p)

    p = doc.add_paragraph()
    p.add_run(
        "实验数据配置文件为“煤块和石块/data.yaml”，其中训练图像目录为 train/images，验证图像目录为 valid/images，"
        "检测类别数为 2，类别名称分别为 coal 和 gangue。训练结果统一保存至 runs/train/YOLO-GDL，"
        "包括 best.pt 权重文件、results.csv、PR 曲线、F1 曲线和混淆矩阵等文件。"
    )
    elements.append(p._p)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run("代码 1 YOLO-GDL 模型训练代码").bold = True
    elements.append(cap._p)

    for line in TRAIN_CODE.splitlines():
        cp = doc.add_paragraph()
        cp.paragraph_format.left_indent = Inches(0.15)
        cp.paragraph_format.right_indent = Inches(0.15)
        cp.paragraph_format.space_after = Pt(0)
        cp.paragraph_format.line_spacing = 1.0
        run = cp.add_run(line if line else " ")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(8.5)
        elements.append(cp._p)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run("表 1 YOLO-GDL 训练参数设置").bold = True
    elements.append(cap._p)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ("参数项", "取值", "说明")
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, "D9EAF7")
        set_cell_text(cell, header, bold=True, size=9)

    for item, value, note in PARAM_ROWS:
        cells = table.add_row().cells
        set_cell_text(cells[0], item, size=8)
        set_cell_text(cells[1], value, size=8)
        set_cell_text(cells[2], note, size=8)

    for row in table.rows:
        row.cells[0].width = Inches(1.35)
        row.cells[1].width = Inches(2.25)
        row.cells[2].width = Inches(3.2)

    elements.append(table._tbl)

    p = doc.add_paragraph()
    p.add_run(
        "上述参数中，imgsz 设置为 1280 是为了保留煤块与矸石表面纹理差异；"
        "hsv_v、hsv_s 等光照增强参数用于模拟井下或输送带现场亮度变化；"
        "close_mosaic 设置为 10，使模型在训练末期更多学习真实图像分布。"
    )
    elements.append(p._p)

    move_after(anchor, elements)

    doc.save(PAPER)
    print(PAPER)


if __name__ == "__main__":
    main()
